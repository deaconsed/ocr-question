"""Export a single subject within an exam session to a Word (.docx) document.

Question text may contain Markdown plus LaTeX math (delimiters: $...$, $$...$$,
\\(...\\), \\[...\\]). Math is rendered to crisp inline images via matplotlib's
mathtext engine (no LaTeX/pandoc system install required). Question diagrams are
embedded from disk.
"""

import os
import re
import json
from io import BytesIO

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.path import Path as MplPath
from matplotlib.patches import PathPatch
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from db import get_db_connection

INPUT_DIR = "extracted_frames"
MATH_DPI = 200
MATH_FONTSIZE = 12
MAX_IMAGE_WIDTH_IN = 4.5

# Math segments in any of the supported delimiter styles. Display ($$, \[) and
# inline ($, \() are all rendered the same way (inline image) for simplicity.
_MATH_RE = re.compile(r'(\$\$.*?\$\$|\\\[.*?\\\]|\$[^$]+?\$|\\\(.*?\\\))', re.S)
# Only bold is auto-detected; single '*' is left alone to avoid mangling
# multiplication signs that appear in plain question text.
_BOLD_RE = re.compile(r'(\*\*.+?\*\*|__.+?__)')

# matplotlib mathtext cannot render LaTeX matrix environments, so we detect them
# and draw them ourselves.
_MATRIX_RE = re.compile(
    r'\\begin\{(bmatrix|Bmatrix|pmatrix|vmatrix|matrix)\}(.*?)\\end\{\1\}', re.S)

# LaTeX backslash-escapes that show up in plain (non-math) text and should be
# turned back into their literal characters in the document.
_LATEX_ESCAPES = [
    ("\\%", "%"), ("\\{", "{"), ("\\}", "}"), ("\\$", "$"),
    ("\\&", "&"), ("\\#", "#"), ("\\_", "_"), ("\\,", " "), ("\\;", " "),
]


def _unescape_latex(text):
    """Turn LaTeX char-escapes (\\%, \\{, \\}, …) into literal characters."""
    for a, b in _LATEX_ESCAPES:
        text = text.replace(a, b)
    return text


def _render_math_png(latex):
    """Render a LaTeX math string to a transparent PNG using matplotlib mathtext.
    Returns (BytesIO, width_in, height_in) or None if the formula can't be parsed."""
    try:
        fig = plt.figure()
        fig.text(0, 0, f"${latex}$", fontsize=MATH_FONTSIZE)
        buf = BytesIO()
        fig.savefig(buf, dpi=MATH_DPI, format="png",
                    bbox_inches="tight", pad_inches=0.02, transparent=True)
        plt.close(fig)
        buf.seek(0)
        with Image.open(buf) as im:
            w_px, h_px = im.size
        buf.seek(0)
        return buf, w_px / MATH_DPI, h_px / MATH_DPI
    except Exception:
        plt.close("all")
        return None


def _split_matrix(latex):
    """Split a math fragment into ('text', str) and ('matrix', (kind, body)) parts
    so any surrounding expression (e.g. 'A = ') is kept around the matrix."""
    parts, pos = [], 0
    for m in _MATRIX_RE.finditer(latex):
        if m.start() > pos:
            parts.append(("text", latex[pos:m.start()]))
        parts.append(("matrix", (m.group(1), m.group(2))))
        pos = m.end()
    if pos < len(latex):
        parts.append(("text", latex[pos:]))
    return parts


def _render_matrix_png(body, kind):
    """Render a LaTeX matrix environment to a PNG with proper delimiters, since
    matplotlib mathtext can't. Returns (BytesIO, width_in, height_in) or None."""
    try:
        rows = [r for r in re.split(r'\\\\', body) if r.strip() != ""]
        grid = [[c.strip() for c in row.split("&")] for row in rows]
        if not grid:
            return None
        ncols = max(len(r) for r in grid)
        nrows = len(grid)
        col_w, row_h = 1.3, 1.1

        fig = plt.figure()
        ax = fig.add_axes([0, 0, 1, 1])
        ax.axis("off")
        for i, row in enumerate(grid):
            for j, cell in enumerate(row):
                if cell == "":
                    continue
                x = j * col_w
                y = (nrows - 1 - i) * row_h
                ax.text(x, y, f"${cell}$", ha="center", va="center",
                        fontsize=MATH_FONTSIZE)

        right_x = (ncols - 1) * col_w
        top_y = (nrows - 1) * row_h + 0.55
        bot_y = -0.55
        left = -0.6
        right = right_x + 0.6
        ax.set_xlim(left - 0.3, right + 0.3)
        ax.set_ylim(bot_y - 0.2, top_y + 0.2)

        lw = 1.4
        serif = 0.18
        if kind in ("bmatrix", "Bmatrix"):
            for xv, s in ((left, serif), (right, -serif)):
                ax.plot([xv, xv], [bot_y, top_y], color="black", lw=lw)
                ax.plot([xv, xv + s], [top_y, top_y], color="black", lw=lw)
                ax.plot([xv, xv + s], [bot_y, bot_y], color="black", lw=lw)
        elif kind == "vmatrix":
            ax.plot([left, left], [bot_y, top_y], color="black", lw=lw)
            ax.plot([right, right], [bot_y, top_y], color="black", lw=lw)
        elif kind == "pmatrix":
            for xv, opening in ((left, 1), (right, -1)):
                ctrl_x = xv - opening * 0.4
                path = MplPath(
                    [(xv, top_y), (ctrl_x, (top_y + bot_y) / 2), (xv, bot_y)],
                    [MplPath.MOVETO, MplPath.CURVE3, MplPath.CURVE3])
                ax.add_patch(PathPatch(path, fill=False, lw=lw))
        # kind == "matrix": no delimiters

        buf = BytesIO()
        fig.savefig(buf, dpi=MATH_DPI, format="png",
                    bbox_inches="tight", pad_inches=0.05, transparent=True)
        plt.close(fig)
        buf.seek(0)
        with Image.open(buf) as im:
            w_px, h_px = im.size
        buf.seek(0)
        return buf, w_px / MATH_DPI, h_px / MATH_DPI
    except Exception:
        plt.close("all")
        return None


def _split_math(text):
    """Yield ('text', str) / ('math', str) segments in document order."""
    segs, pos = [], 0
    for m in _MATH_RE.finditer(text):
        if m.start() > pos:
            segs.append(("text", text[pos:m.start()]))
        tok = m.group(1)
        if tok.startswith("$$") and tok.endswith("$$"):
            inner = tok[2:-2]
        elif tok.startswith(r"\[") and tok.endswith(r"\]"):
            inner = tok[2:-2]
        elif tok.startswith(r"\(") and tok.endswith(r"\)"):
            inner = tok[2:-2]
        else:
            inner = tok[1:-1]
        segs.append(("math", inner.strip()))
        pos = m.end()
    if pos < len(text):
        segs.append(("text", text[pos:]))
    return segs


def _add_text_with_bold(paragraph, text):
    """Add plain text to a paragraph, honouring **bold**/__bold__ markers and
    unescaping LaTeX char-escapes (\\%, \\{, …) so they don't print literally."""
    if not text:
        return
    for part in _BOLD_RE.split(text):
        if not part:
            continue
        if (part.startswith("**") and part.endswith("**")) or \
           (part.startswith("__") and part.endswith("__")):
            paragraph.add_run(_unescape_latex(part[2:-2])).bold = True
        else:
            paragraph.add_run(_unescape_latex(part))


def _add_picture_run(paragraph, res):
    buf, w_in, _ = res
    paragraph.add_run().add_picture(buf, width=Inches(w_in))


def _render_math_segment(paragraph, latex):
    """Render one math fragment inline. Handles matrix environments (which
    mathtext can't) by drawing them and rendering any surrounding expression."""
    if not _MATRIX_RE.search(latex):
        res = _render_math_png(latex)
        if res:
            _add_picture_run(paragraph, res)
        else:
            paragraph.add_run(f"${_unescape_latex(latex)}$")
        return

    for kind, payload in _split_matrix(latex):
        if kind == "matrix":
            mkind, body = payload
            res = _render_matrix_png(body, mkind)
            if res:
                _add_picture_run(paragraph, res)
            else:
                paragraph.add_run(_unescape_latex(body))
        else:
            if not payload.strip():
                continue
            res = _render_math_png(payload)
            if res:
                _add_picture_run(paragraph, res)
            else:
                paragraph.add_run(_unescape_latex(payload))


def _add_rich_text(paragraph, text):
    """Add inline text with math + light markdown into one existing paragraph
    (used for options, which are single-line)."""
    for kind, content in _split_math(text or ""):
        if kind == "text":
            _add_text_with_bold(paragraph, content)
        else:
            _render_math_segment(paragraph, content)


def _add_question_body(doc, text):
    """Add question text that may span multiple lines and contain inline math
    (math fragments may themselves span newlines, e.g. a matrix)."""
    para = doc.add_paragraph()
    for kind, content in _split_math(text):
        if kind == "text":
            lines = content.split("\n")
            for i, line in enumerate(lines):
                if i > 0:
                    para = doc.add_paragraph()
                if line:
                    _add_text_with_bold(para, line)
        else:
            _render_math_segment(para, content)


def _question_image_path(exam_id, subject, filename):
    subj_folder = subject.lower().replace(" ", "_")
    primary = os.path.join(INPUT_DIR, f"exam_{exam_id}", subj_folder, filename)
    if os.path.isfile(primary):
        return primary
    legacy = os.path.join(INPUT_DIR, subj_folder, filename)
    return legacy if os.path.isfile(legacy) else None


def _add_image(doc, path):
    try:
        with Image.open(path) as im:
            w_px = im.size[0]
            dpi = (im.info.get("dpi") or (96, 96))[0] or 96
        width_in = min(MAX_IMAGE_WIDTH_IN, max(1.0, w_px / dpi))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width_in))
    except Exception:
        pass


def _safe_filename(name):
    return re.sub(r'[^A-Za-z0-9._-]+', '_', name).strip('_') or "export"


def build_subject_docx(exam_id, subject):
    """Build the .docx for one subject in one session.
    Returns (BytesIO, suggested_filename). Raises ValueError if exam not found."""
    conn = get_db_connection()
    if not conn:
        raise RuntimeError("Database connection failed")
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute(
            "SELECT id, year, label, session_index FROM exams WHERE id = %s",
            (exam_id,))
        exam = cursor.fetchone()
        if not exam:
            raise ValueError("Exam session not found")
        cursor.execute("""
            SELECT q.question_number, q.question_text, q.options,
                   q.question_image, q.teacher_answer
            FROM questions q
            JOIN subjects s ON q.subject_id = s.id
            WHERE s.name = %s AND q.exam_id = %s
            ORDER BY q.question_number ASC
        """, (subject, exam_id))
        questions = cursor.fetchall()
    finally:
        cursor.close()
        conn.close()

    doc = Document()

    subtitle_bits = []
    if exam.get("year"):
        subtitle_bits.append(str(exam["year"]))
    subtitle_bits.append(exam.get("label") or f"Session {exam.get('session_index')}")
    subtitle = " · ".join(subtitle_bits)

    doc.add_heading(subject.upper(), level=0)
    sub = doc.add_paragraph()
    sub_run = sub.add_run(subtitle)
    sub_run.italic = True
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if not questions:
        doc.add_paragraph("No questions found for this subject in this session.")
    else:
        for q in questions:
            head = doc.add_paragraph()
            hr = head.add_run(f"Question {q['question_number']}")
            hr.bold = True
            hr.font.size = Pt(13)

            text = q.get("question_text")
            if text:
                _add_question_body(doc, str(text))

            if q.get("question_image"):
                path = _question_image_path(exam_id, subject, q["question_image"])
                if path:
                    _add_image(doc, path)

            options = {}
            if q.get("options"):
                try:
                    options = json.loads(q["options"]) if isinstance(q["options"], str) else q["options"]
                except Exception:
                    options = {}
            for letter in ("A", "B", "C", "D"):
                if options.get(letter):
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Inches(0.3)
                    para.add_run(f"{letter}.  ").bold = True
                    _add_rich_text(para, str(options[letter]))

            if q.get("teacher_answer"):
                ans = doc.add_paragraph()
                ar = ans.add_run(f"Answer: {q['teacher_answer']}")
                ar.bold = True
                ar.font.color.rgb = RGBColor(0x1a, 0x7f, 0x37)

            doc.add_paragraph()  # spacer between questions

    out = BytesIO()
    doc.save(out)
    out.seek(0)

    label_part = exam.get("label") or f"session{exam.get('session_index')}"
    filename = _safe_filename(f"{subject}_{label_part}") + ".docx"
    return out, filename
